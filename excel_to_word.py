## أحرص على تعديل الكود فيما يتناسق مع ملفاتك

import pandas as pd  # Import pandas for working with Excel data
from docx import Document  # Import python-docx to create and edit Word documents
import os  # Import os for file path and folder operations

# Load the Excel file and read the first 152 rows from the "COMPUTER ISSUE" sheet
# تحميل ملف Excel مع تحديد ورقة العمل "COMPUTER ISSUE" وقراءة أول 152 صفًا فقط
excel_file = r"path_to_your_excel_file.xlsx"  # استبدل بمسار ملف Excel
df = pd.read_excel(excel_file, sheet_name='COMPUTER ISSUE')  # Specify the required sheet
df = df.iloc[:152]  # Limit to the first 152 rows (adjustable based on needs)

# Output folder path for saving the Word files
# مسار مجلد الإخراج لحفظ ملفات Word الناتجة
output_folder = r"path_to_output_folder"  # استبدل بمسار مجلد الإخراج
if not os.path.exists(output_folder):  # Check if the folder exists
    os.makedirs(output_folder)  # Create the folder if it doesn't exist

# Function to extract processor type and manufacturer
# فلترة لعامود في الاكسل لاستخراج بيانات معينه فقط بدلا من اخذها كلها 
def extract_processor_and_manufacturer(processor_text, manufacturer_text):
    processor_type = ''  # Default to empty if neither i5 nor i7 is found
    
    # Ensure processor type contains i5 or i7
    # هنا سوينا فلتر انه فقط يستخرج لي اختصار المعالج بدلا من اخذ اسمه كامل 
    if isinstance(processor_text, str):  # Check if the text is a string
        if 'i5' in processor_text:
            processor_type = 'i5'
        elif 'i7' in processor_text:
            processor_type = 'i7'

    # Combine manufacturer and processor type if available
    # دمج اسم الشركة مع نوع المعالج إذا وجد
    if isinstance(manufacturer_text, str):  # Check if the manufacturer is a string
        if processor_type:
            return f"{manufacturer_text} - {processor_type} core"
        else:
            return ""  # Leave blank if neither i5 nor i7
    return ""  # Default return in case of invalid inputs

# Function to create a Word document for each employee
# دالة لإنشاء ملف Word لكل موظف بناءً على بيانات ملف Excel
def generate_word_doc(employee):
    # Load the Word template
    # تحميل قالب Word
    template = r"path_to_word_template.docx"  # استبدل بمسار قالب Word
    try:
        doc = Document(template)  # Load the Word template
    except Exception as e:  # Handle errors in loading the template
        print(f"Error loading template: {e}")  # طباعة رسالة الخطأ عند فشل تحميل القالب
        return
    
    # Replace placeholders in the Word document
    # استبدال النصوص المؤقتة في قالب Word
    for paragraph in doc.paragraphs:
        if '[Ansar Abbas]' in paragraph.text:  # Replace employee name placeholder
            paragraph.text = paragraph.text.replace('[Ansar Abbas]', employee['Name'])
        if '[Date14/11/2024]' in paragraph.text:  # Replace date placeholder
            paragraph.text = paragraph.text.replace('[Date14/11/2024]', pd.Timestamp.today().strftime('%d/%m/%Y'))
    
    # Update the table in the template
    # تحديث الجدول في القالب
    table = doc.tables[0]
    
    # Extract processor and manufacturer details
    # استخراج تفاصيل المعالج والمصنع
    processor_and_manufacturer = extract_processor_and_manufacturer(employee['Processor'], employee['System Manufacturer'])
    
    # Populate the table with employee details
    # تعبئة الجدول ببيانات الموظف
    table.cell(2, 2).text = str(employee['Machine name'])  # Machine name
    table.cell(2, 3).text = str(employee['CPU Serial No'])  # CPU Serial No
    table.cell(2, 4).text = str(employee['Name'])  # Employee name
    table.cell(2, 5).text = processor_and_manufacturer  # Processor and Manufacturer details
    
    # Fill in LCD data in the fourth row
    # تعبئة بيانات LCD في الصف الرابع
    table.cell(3, 2).text = str(employee['LCD TAG'])  # LCD TAG
    table.cell(3, 3).text = str(employee['LCD Serial No'])  # LCD Serial No
    table.cell(3, 4).text = str(employee['Name'])  # User name
    table.cell(3, 5).text = ""  # Leave the last column blank
    
    # Define the output file name and save the document
    # تسمية كل ملف بنفس اسم الموظف
    output_filename = os.path.join(output_folder, f"{employee['Name']}_handover.docx")  # Output file based on employee name
    doc.save(output_filename)  # Save the file
    print(f"File created: {output_filename}")  # Print confirmation message

# Loop through all employees in the Excel file and create Word documents
# تكرار العملية لكل موظف في ملف Excel
for _, row in df.iterrows():
    generate_word_doc(row)  # Call the function to create Word files
