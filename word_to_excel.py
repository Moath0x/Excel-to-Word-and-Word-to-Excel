import os
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

# مسار المجلد الذي يحتوي على ملفات Word - استبدل "ادخل مسار الملف" بمسار مجلدك
# Path to the folder containing Word files - Replace "Enter folder path" with your folder path
input_folder = r"Enter folder path"

# مسار ملف Excel الناتج - استبدل "ادخل مسار الملف" بمسار ملف Excel المطلوب
# Path to the output Excel file - Replace "Enter file path" with your desired Excel file path
output_excel = r"Enter file path"

# قائمة لتخزين البيانات المستخرجة
# List to store extracted data
data = []

# دالة لقراءة البيانات من ملف Word
# Function to read data from a Word file
def extract_data_from_word(file_path):
    """
    استخراج البيانات من ملف Word معين.
    Extract data from a specific Word file.
    Args:
        file_path (str): مسار ملف Word. / Path to the Word file.
    Returns:
        dict: بيانات مستخرجة كقاموس. / Extracted data as a dictionary.
    """
    try:
        doc = Document(file_path)  # تحميل ملف Word / Load the Word file
    except Exception as e:
        print(f"خطأ أثناء قراءة الملف {file_path}: {e}")  # Error reading file
        return None

    # قائمة لتخزين النصوص المستخرجة
    # Dictionary to store extracted data
    extracted_data = {}

    # قراءة الجدول الأول في الملف (إذا كان موجودًا)
    # Read the first table in the file (if it exists)
    if doc.tables:
        table = doc.tables[0]
        try:
            extracted_data['User'] = table.cell(2, 4).text.strip()  # استخراج اسم المستخدم / Extract user name
            extracted_data['Machine Name'] = table.cell(2, 2).text.strip()
            extracted_data['CPU Serial No'] = table.cell(2, 3).text.strip()
            extracted_data['LCD TAG'] = table.cell(3, 2).text.strip()
            extracted_data['LCD Serial No'] = table.cell(3, 3).text.strip()
        except IndexError as e:
            print(f"خطأ أثناء قراءة الجدول في الملف {file_path}: {e}")  # Error reading table in file

    return extracted_data

# قراءة جميع الملفات في المجلد
# Loop through all files in the folder
for filename in os.listdir(input_folder):
    if filename.endswith('.docx'):  # التأكد من أن الملف هو ملف Word / Ensure the file is a Word file
        file_path = os.path.join(input_folder, filename)
        data_entry = extract_data_from_word(file_path)
        if data_entry:  # إذا تم استخراج البيانات بنجاح / If data was successfully extracted
            data.append(data_entry)

# إنشاء ملف Excel باستخدام Pandas
# Create an Excel file using Pandas
if data:
    df = pd.DataFrame(data)  # تحويل البيانات إلى DataFrame / Convert data to a DataFrame
    
    # تعديل ترتيب الأعمدة / Rearrange columns
    desired_columns = ['User', 'Machine Name', 'CPU Serial No', 'LCD TAG', 'LCD Serial No']
    df = df[desired_columns]
    
    df.to_excel(output_excel, index=False)  # حفظ البيانات في ملف Excel / Save data to an Excel file
    print(f"تم إنشاء ملف Excel: {output_excel}")  # Excel file created

    # تحسين تنسيق الأعمدة باستخدام openpyxl
    # Enhance column formatting using openpyxl
    wb = load_workbook(output_excel)
    ws = wb.active
    for col in ws.columns:
        column_letter = get_column_letter(col[0].column)
        ws.column_dimensions[column_letter].width = 20  # تعيين عرض ثابت لكل عامود / Set a fixed width for each column
    wb.save(output_excel)
    print("DONE!")  # Task completed
else:
    print("ERROR")  # No data was extracted
